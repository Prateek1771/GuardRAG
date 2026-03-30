import io
import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from typing import List


def parse_document(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_bytes)
    elif ext in ("xlsx", "xls", "csv"):
        return _parse_excel(file_bytes, ext)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        return file_bytes.decode("utf-8", errors="ignore")


def _parse_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
            # Also extract tables
            for table in page.extract_tables():
                for row in table:
                    cleaned = [str(cell or "").strip() for cell in row]
                    text_parts.append(" | ".join(cleaned))
    return "\n\n".join(text_parts)


def _parse_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _parse_excel(file_bytes: bytes, ext: str) -> str:
    parts = []
    if ext == "csv":
        df_map = {"Sheet1": pd.read_csv(io.BytesIO(file_bytes))}
    else:
        df_map = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, engine="openpyxl")

    for sheet_name, df in df_map.items():
        parts.append(f"## Sheet: {sheet_name}")
        # Preserve table structure as readable text
        parts.append(df.to_string(index=True))
        parts.append("")  # blank line between sheets
    return "\n".join(parts)
